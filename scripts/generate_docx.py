from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Add Title
title = doc.add_heading('EduTech Reels - Comprehensive Technical Q&A (Compiled Specs)', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p_intro = doc.add_paragraph("This document compiles the core technical specifications across Database Schema, Deployment, Performance, Recommendation Algorithms, and Scalability Infrastructure into a unified Q&A format. Redundancies have been eliminated to provide a streamlined, in-depth view for enterprise technical teams.")
p_intro.italic = True

def add_qa(q, a):
    p = doc.add_paragraph()
    run_q = p.add_run(q)
    run_q.bold = True
    p_a = doc.add_paragraph(a)

doc.add_heading('Section 1: Database Schema & Data Strategy', level=1)
add_qa("Q1: What is the underlying database schema used for the recommendation engine, and why was Firestore chosen?",
       "A: We use a highly denormalized NoSQL schema optimized for reads using Firebase Firestore. Firestore was chosen due to its ability to handle massive read scalability and sub-10ms response times. The core collections are `users` (~1.5 KB/doc), `videos` (~2 KB/doc), `interactions` (Engagement Logs), `creators`, and separated collections for `user_embeddings` & `video_embeddings` to prevent massive payload sizes during standard reads.")

add_qa("Q2: How are real-time feeds generated quickly without bottlenecking the database?",
       "A: We leverage Redis (Cloud Memorystore) as a caching layer. The API first checks Redis for cached user feeds or video IDs. On a cache miss, it fetches embeddings, calls Vertex AI for recommendations, and caches the result. Fetching video metadata is done via an 'in' query against Firestore using the batched IDs, ensuring fast O(1) lookups.")

add_qa("Q3: How does the system plan to migrate when Firestore's write limits become a bottleneck?",
       "A: Firestore has a hard limit of 10,000 writes/sec per database. At ~500K to 1M DAU, the relational video catalog and engagement logs migrate to Google Cloud Spanner. Spanner uses TrueTime technology to provide strict external consistency globally while allowing unbounded horizontal scaling through primary key sharding (e.g., prefixing keys with UUIDs or hashes to prevent hotspots).")

add_qa("Q4: How does the platform handle Backup and Disaster Recovery for massive datasets?",
       "A: Point-in-Time Recovery (PITR) is enabled in Firestore to revert to any microsecond in the past 7 days instantly. Furthermore, we run nightly automated exports via Cloud Scheduler to Cloud Storage (Coldline) with a 1-year retention policy for ML training and compliance.")

doc.add_heading('Section 2: Recommendation Algorithm & Mathematics', level=1)
add_qa("Q5: How does the hybrid recommendation algorithm mathematically calculate the Final Ranking Score for a video?",
       "A: The algorithm computes a Final Ranking Score (S_final) as a linear combination of three signals: S_final = (0.5 * S_content) + (0.3 * S_collab) + (0.2 * S_trending). This weights explicit educational relevance highest, followed by peer-validated discovery, and finally fresh/viral content.")

add_qa("Q6: How is the Content-Based Filtering Score (S_content) computed?",
       "A: S_content measures the direct affinity between user interests and video metadata. It aggregates three sub-scores: Category Match (1.0 if aspirant types match), Tag Overlap (Jaccard similarity between user and video tags normalized by user tag count), and Creator Affinity (based on a weighted sum of historical views and likes for the video's creator).")

add_qa("Q7: How does the system implement Collaborative Filtering (S_collab)?",
       "A: We use a User-User KNN (K-Nearest Neighbors) similarity approach. We find the top K users with the highest Jaccard similarity to the current user based on their liked/saved video history. The collaborative score for an unseen video is the weighted average of the engagements of these similar users.")

add_qa("Q8: How is the Popularity/Trending Score (S_trending) calculated to avoid surfacing only old, viral videos?",
       "A: The Trending score uses a weighted sum of raw engagements. To prioritize fresh content, an exponential decay function is applied based on the video's age in hours. This ensures that new videos quickly replace older viral content on the global trending list.")

doc.add_heading('Section 3: Scalability & Infrastructure Architecture', level=1)
add_qa("Q9: Can you detail the production infrastructure stack required to handle 1M+ DAU?",
       "A: The architecture uses GCP: Cloud CDN serves static assets and video chunks at edge nodes. A Global Load Balancer routes API traffic to Google Kubernetes Engine (GKE) clusters running stateless microservices. Redis caches hot data, Firestore/Spanner serves as the database, Cloud Storage holds videos, and Vertex AI handles real-time ML inference.")

add_qa("Q10: What is the load handling strategy during peak concurrency (e.g., 50,000 concurrent users)?",
       "A: At peak, video bandwidth approaches 80 Gbps, which is entirely offloaded to Cloud CDN. API requests scale horizontally via GKE's Horizontal Pod Autoscaler (HPA) when CPU > 65%. Heavy computation (like recommendation sorting) is shifted to background workers and Redis Sorted Sets to ensure API pods only perform lightweight data fetching.")

doc.add_heading('Section 4: Performance Metrics & Evaluation', level=1)
add_qa("Q11: What metrics are used to evaluate the algorithm offline before deploying?",
       "A: We use Precision@K, Recall@K, NDCG (Normalized Discounted Cumulative Gain to heavily penalize algorithms that place relevant videos at the bottom), Hit Rate, and Mean Average Precision (MAP). The targets are Precision@10 > 0.40 and NDCG@10 > 0.65.")

add_qa("Q12: How do you measure online performance and user satisfaction?",
       "A: We use A/B testing (e.g., 10% Control, 45% Variant A, 45% Variant B) and measure Watch-Through Rate (target > 60%), Engagement Rate (target > 8%), Session Duration (target > 12 mins), and D7 Retention. Moving from random recommendations to the Hybrid model increases Session Duration by approximately 6x.")

doc.add_heading('Section 5: Deployment, MLOps, & Security', level=1)
add_qa("Q13: How are code and ML model updates deployed with zero downtime?",
       "A: We use a Monorepo with a CI/CD pipeline orchestrated via GitHub Actions to deploy to GKE using a Canary rollout strategy. If 5xx error rates spike >1% during the rollout pause window, GKE automatically rolls back.")

add_qa("Q14: How does the MLOps pipeline handle model decay or 'Concept Drift'?",
       "A: Models are deployed to Vertex AI Endpoints. If offline validation metrics drop or user engagement dips, Vertex AI Model Monitoring triggers a 'Concept Drift' alert. This automatically launches a Vertex AI Pipeline to retrain the model on the latest 30 days of interaction data stored in BigQuery.")

add_qa("Q15: What security measures are in place to protect the platform and user data?",
       "A: Compute instances reside in private VPC subnets with no public IPs. Secrets are injected via Google Secret Manager. Cloud Armor mitigates DDoS attacks at the load balancer level. API authentication utilizes short-lived JWTs, and rigorous vulnerability scanning is performed via Google Artifact Registry.")

file_path = 'technical documentation/Compiled_Technical_Documentation.docx'
doc.save(file_path)
print(f"Document successfully saved to {file_path}")
