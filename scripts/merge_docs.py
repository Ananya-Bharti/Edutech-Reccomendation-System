import os
from docx import Document
from docxcompose.composer import Composer

directory = 'technical documentation'
files = [
    'Database_Schema.docx',
    'Recommendation_Algorithm.docx',
    'Scalability_Infrastructure.docx',
    'Deployment_Operations.docx',
    'Performance_Analysis.docx'
]

# Create a blank master document
master = Document()
master.add_heading('EduTech Reels - Final Technical Documentation', 0)
composer = Composer(master)

for f in files:
    filepath = os.path.join(directory, f)
    if os.path.exists(filepath):
        print(f"Appending {f}...")
        doc = Document(filepath)
        composer.append(doc)

output_path = os.path.join(directory, 'Final_Technical_Documentation.docx')
composer.save(output_path)
print(f"Successfully compiled into {output_path}")
