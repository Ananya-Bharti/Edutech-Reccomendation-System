from docx import Document
from docx.shared import Pt
import os

doc = Document()
doc.add_heading('EduTech Reels - Frontend Architecture & Scoring System', 0)

doc.add_heading('1. Frontend Architecture Overview', level=1)
doc.add_paragraph(
    "The frontend is built using Next.js (React) with a heavily optimized client-side state management system to ensure zero-latency feed re-ranking during the demo. "
    "The application relies on React hooks (useState, useMemo, useCallback, useRef) to maintain an in-memory Engagement Engine that tracks user interactions without waiting for backend network requests."
)

p1 = doc.add_paragraph()
p1.add_run("Key Components:\n").bold = True
p1.add_run("- Engagement Engine: ").bold = True
p1.add_run("Maintains a per-user state containing video interaction records (views, likes, saves, watch time) and accumulated tag/creator affinity scores.\n")
p1.add_run("- Lazy Video Tile: ").bold = True
p1.add_run("Implements an IntersectionObserver. Videos are only rendered and auto-played when they scroll into the viewport, saving bandwidth and accurately triggering watch-time tracking.\n")
p1.add_run("- Real-time Feed Construction: ").bold = True
p1.add_run("Whenever a user interacts with a video (e.g., clicks 'Like'), the engagement state mutates immutably, triggering a recalculation of the user's affinity scores and an immediate re-sort of the feed grid.")

doc.add_heading('2. Engagement Interaction Weights', level=1)
doc.add_paragraph("Every user interaction is tracked and assigned a base weight. These weights determine how rapidly the system learns the user's preferences:")
p2 = doc.add_paragraph()
p2.add_run("- View (opening the modal/viewport): ").bold = True; p2.add_run("1 point\n")
p2.add_run("- Like: ").bold = True; p2.add_run("3 points\n")
p2.add_run("- Save/Bookmark: ").bold = True; p2.add_run("5 points\n")
p2.add_run("- Share: ").bold = True; p2.add_run("5 points\n")
p2.add_run("- Watch Time: ").bold = True; p2.add_run("2 points per 10 seconds watched")

doc.add_paragraph("These raw engagement points are distributed to the video's underlying tags and its creator. For example, if a user generates 10 points on a video tagged 'finance', the user's affinity for 'finance' increases by 10.")

doc.add_heading('3. The Scoring Algorithm (scoreVideo)', level=1)
doc.add_paragraph("For every video in the database, the frontend calculates a dynamic ranking score based on six distinct factors:")

p3 = doc.add_paragraph()
p3.add_run("Factor 1: Category Match (Base Score)\n").bold = True
p3.add_run("- Undecided Users (Blank Slate): +20 points for all videos, ensuring a level playing field.\n")
p3.add_run("- Typed Users: +40 points if the video matches their primary aspirant type (e.g., MBA for an MBA user), and +8 points for cross-category videos.\n\n")

p3.add_run("Factor 2: Static Interest Match\n").bold = True
p3.add_run("- +8 points for every video tag that exactly matches the user's declared profile interests.\n")
p3.add_run("- +6 points if any of the user's interest keywords appear directly in the video title.\n\n")

p3.add_run("Factor 3: Learned Tag Affinity (Amplified)\n").bold = True
p3.add_run("For videos the user has NOT interacted with yet:\n")
p3.add_run("- The algorithm adds (Tag Engagement Score × 3), capped at a maximum of +40 points per tag. This is the primary driver for dynamic feed reshuffling.\n\n")

p3.add_run("Factor 4: Learned Creator Affinity\n").bold = True
p3.add_run("For videos the user has NOT interacted with yet:\n")
p3.add_run("- The algorithm adds (Creator Engagement Score × 3), capped at a maximum of +35 points.\n\n")

p3.add_run("Factor 5: Already-Seen Penalty (Freshness Booster)\n").bold = True
p3.add_run("To ensure users discover new content rather than seeing the same highly-scored videos over and over, interacted videos are heavily penalized:\n")
p3.add_run("- Base Penalty for viewing: -40 points.\n")
p3.add_run("- Liked Penalty: -15 points.\n")
p3.add_run("- Saved Penalty: -10 points.\n")
p3.add_run("- Shared Penalty: -15 points.\n\n")

p3.add_run("Factor 6: Random Jitter\n").bold = True
p3.add_run("- Adds +0 to +3 random points to break ties and provide slight variations during every re-render.")

doc.add_heading('4. Feed Interleaving Strategies', level=1)
doc.add_paragraph("Once all videos are scored, the feed is constructed using one of two interleaving methods:")
p4 = doc.add_paragraph()
p4.add_run("- Undecided User Strategy: ").bold = True
p4.add_run("Videos are sorted purely by descending score. However, a 'Diversity Cap' is applied: the system ensures that no more than 3 consecutive videos from the same category appear in a row.\n")
p4.add_run("- Typed User Strategy: ").bold = True
p4.add_run("The system interleaves the videos using a 4:1 ratio. It displays 4 top-scoring videos from the user's primary category, followed by 1 top-scoring cross-category video to encourage discovery.")

os.makedirs('technical documentation', exist_ok=True)
doc.save('technical documentation/frontend and scoring.docx')
print("Successfully generated frontend and scoring.docx")
